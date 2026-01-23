#!/usr/bin/env python3
"""
Convert .md to DOCX with embedded diagrams.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

DOCS_DIR = Path(__file__).parent.parent / "docs"
MD_FILE = DOCS_DIR / "dremio-architecture.md"
OUTPUT_FILE = DOCS_DIR / "dremio-architecture.docx"

# Map diagram references to image files
DIAGRAM_MAP = {
    "k8s-overview-diagram.mermaid": DOCS_DIR / "k8s-overview-diagram.png",
    "k8s-detailed-diagram.mermaid": DOCS_DIR / "k8s-detailed-diagram.png",
    "storage-diagram.mermaid": DOCS_DIR / "storage-diagram.png",
}


def setup_styles(doc):
    """Configure document styles."""
    styles = doc.styles

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2d, 0x3e, 0x50)

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.size = Pt(12)
    h3.font.bold = True

    # Normal
    normal = styles['Normal']
    normal.font.size = Pt(11)
    normal.font.name = 'Calibri'


def add_table(doc, headers, rows):
    """Add a table to the document."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            if i < len(row_cells):
                row_cells[i].text = cell_text

    doc.add_paragraph()


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx."""
    headers = []
    rows = []

    # Parse header
    header_line = lines[start_idx].strip()
    if '|' in header_line:
        headers = [h.strip() for h in header_line.split('|')[1:-1]]

    # Skip separator line
    i = start_idx + 2

    # Parse rows
    while i < len(lines) and '|' in lines[i]:
        row_line = lines[i].strip()
        if row_line.startswith('|'):
            row_data = [c.strip() for c in row_line.split('|')[1:-1]]
            rows.append(row_data)
        i += 1

    return headers, rows, i


def process_markdown(doc, md_content):
    """Process markdown content and add to document."""
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_content = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph()
                p.style = 'Normal'
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                doc.add_paragraph()
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            doc.add_paragraph('─' * 50)
            i += 1
            continue

        # Tables
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            headers, rows, next_i = parse_table(lines, i)
            if headers:
                add_table(doc, headers, rows)
            i = next_i
            continue

        # Check for diagram references and embed images
        diagram_match = re.search(r'\[.*?\]\(([\w-]+\.mermaid)\)', line)
        if diagram_match:
            diagram_file = diagram_match.group(1)
            if diagram_file in DIAGRAM_MAP:
                img_path = DIAGRAM_MAP[diagram_file]
                if img_path.exists():
                    # Add caption before image
                    p = doc.add_paragraph()
                    p.add_run(f"Diagram: {diagram_file.replace('.mermaid', '')}").italic = True

                    # Add image
                    doc.add_picture(str(img_path), width=Inches(6.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                    i += 1
                    continue

        # Regular paragraph
        if line.strip():
            p = doc.add_paragraph()

            # Process inline formatting
            text = line.strip()

            # Bold
            parts = re.split(r'\*\*(.+?)\*\*', text)
            for j, part in enumerate(parts):
                if j % 2 == 1:  # Bold text
                    p.add_run(part).bold = True
                else:
                    # Check for inline code
                    code_parts = re.split(r'`(.+?)`', part)
                    for k, code_part in enumerate(code_parts):
                        if k % 2 == 1:  # Code
                            run = p.add_run(code_part)
                            run.font.name = 'Consolas'
                            run.font.size = Pt(10)
                        else:
                            p.add_run(code_part)

        i += 1


def main():
    """Main function to create DOCX."""
    print(f"Reading {MD_FILE}...")
    md_content = MD_FILE.read_text(encoding='utf-8')

    print("Creating DOCX document...")
    doc = Document()

    # Setup styles
    setup_styles(doc)

    # Set document properties
    core_props = doc.core_properties
    core_props.title = "Dremio Enterprise - Kubernetes Architecture"
    core_props.author = "Architecture Team"

    # Process markdown
    process_markdown(doc, md_content)

    # Save document
    print(f"Saving to {OUTPUT_FILE}...")
    doc.save(str(OUTPUT_FILE))
    print(f"Done! Created {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
